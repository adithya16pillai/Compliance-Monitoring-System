import os
import shutil
from typing import List
from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
import PyPDF2
import docx
import aiofiles

from app.database import get_database, create_tables
from app.models import Document, ComplianceResult
from app.compliance_engine import ComplianceEngine

# Create FastAPI app
app = FastAPI(title="Compliance Monitoring API", version="1.0.0")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize compliance engine
compliance_engine = ComplianceEngine()

# Create database tables
create_tables()

# Create uploads directory
os.makedirs("uploads", exist_ok=True)

@app.get("/")
async def root():
    return {"message": "Compliance Monitoring API is running"}

@app.get("/health")
async def health_check():
    return {"status": "healthy", "message": "API is operational"}

def extract_text_from_file(file_path: str, filename: str) -> str:
    """Extract text from uploaded files"""
    text = ""
    file_extension = filename.lower().split('.')[-1]
    
    try:
        if file_extension == 'pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        
        elif file_extension == 'docx':
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        
        elif file_extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text: {str(e)}")
    
    return text

@app.post("/upload-document/")
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_database)
):
    """Upload and process a document"""
    
    # Validate file type
    allowed_extensions = ['pdf', 'docx', 'txt']
    file_extension = file.filename.lower().split('.')[-1]
    
    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"File type {file_extension} not supported. Allowed: {allowed_extensions}"
        )
    
    # Save file
    file_path = f"uploads/{file.filename}"
    
    try:
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        # Extract text
        text_content = extract_text_from_file(file_path, file.filename)
        
        # Save to database
        db_document = Document(
            filename=file.filename,
            file_path=file_path,
            content=text_content,
            file_size=len(content),
            file_type=file_extension
        )
        
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        
        return {
            "message": "File uploaded successfully",
            "document_id": db_document.id,
            "filename": file.filename,
            "file_size": len(content),
            "text_length": len(text_content)
        }
    
    except Exception as e:
        # Clean up file if database save fails
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.post("/analyze-compliance/{document_id}")
async def analyze_compliance(
    document_id: int,
    rule_types: List[str] = None,
    db: Session = Depends(get_database)
):
    """Analyze document for compliance violations"""
    
    # Get document from database
    document = db.query(Document).filter(Document.id == document_id).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # Analyze compliance
        results = compliance_engine.analyze_document(
            doc_id=str(document_id),
            content=document.content,
            rule_types=rule_types
        )
        
        # Save results to database
        for result in results:
            db_result = ComplianceResult(
                document_id=document_id,
                rule_type=result["rule_type"],
                violation_type=result["violation_type"],
                confidence_score=result["confidence_score"],
                evidence=str(result.get("llm_evidence", [])),
                explanation=result["llm_explanation"],
                is_violation=result["overall_violation"]
            )
            db.add(db_result)
        
        db.commit()
        
        # Generate report
        report = compliance_engine.generate_compliance_report(results)
        
        return {
            "document_id": document_id,
            "document_name": document.filename,
            "analysis_results": report
        }
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@app.get("/documents/")
async def list_documents(db: Session = Depends(get_database)):
    """List all uploaded documents"""
    documents = db.query(Document).all()
    
    return {
        "documents": [
            {
                "id": doc.id,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "uploaded_at": doc.uploaded_at
            }
            for doc in documents
        ]
    }

@app.get("/compliance-results/{document_id}")
async def get_compliance_results(
    document_id: int,
    db: Session = Depends(get_database)
):
    """Get compliance results for a document"""
    
    results = db.query(ComplianceResult).filter(
        ComplianceResult.document_id == document_id
    ).all()
    
    if not results:
        raise HTTPException(status_code=404, detail="No compliance results found")
    
    return {
        "document_id": document_id,
        "results": [
            {
                "id": result.id,
                "rule_type": result.rule_type,
                "violation_type": result.violation_type,
                "confidence_score": result.confidence_score,
                "evidence": result.evidence,
                "explanation": result.explanation,
                "is_violation": result.is_violation,
                "created_at": result.created_at
            }
            for result in results
        ]
    }

@app.get("/rules/")
async def get_available_rules():
    """Get all available compliance rules"""
    return {
        "rules": compliance_engine.rules
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)